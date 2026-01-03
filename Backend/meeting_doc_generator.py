import os, json, re
from dotenv import load_dotenv
from docx import Document
from openai import AzureOpenAI

load_dotenv()

# ==========================================
# 환경 변수 불러오기
# ==========================================
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
AZURE_OPENAI_DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")

# 클라이언트 초기화
openai_client = AzureOpenAI(
    api_key=AZURE_OPENAI_API_KEY,
    api_version=AZURE_OPENAI_API_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT
)

# ==========================================
# 1. 스타일 복사 헬퍼 함수
# ==========================================
def update_paragraph_with_style(paragraph, new_text):
    """
    문단(Paragraph)의 기존 스타일을 유지하면서 텍스트만 교체합니다.
    """
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    # 첫 번째 Run의 스타일 복사 및 적용
    paragraph.runs[0].text = new_text
    
    # 잔여 텍스트(Run) 삭제
    for _ in range(len(paragraph.runs) - 1):
        r = paragraph.runs[1]
        r._element.getparent().remove(r._element)

# ==========================================
# 2. 좌표를 포함한 내용 추출 (Template 읽기)
# ==========================================
def extract_text_with_coordinates(docx_path):
    doc = Document(docx_path)
    full_text_list = []
    
    # 문단 추출
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            full_text_list.append(f"[P-{i}] {para.text}")

    # 표 추출
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip().replace("\n", " ")
                if cell_text:
                    full_text_list.append(f"[T-{t_idx}-R-{r_idx}-C-{c_idx}] {cell_text}")
            
    return "\n".join(full_text_list)

# ==========================================
# 3. LLM에 좌표 기반 매핑 요청
# ==========================================
def get_coordinates_json_from_llm(template_text_with_coords, new_summary_text):
    system_prompt = """
    너는 Word 문서 템플릿의 구조를 파악하고 데이터를 교체하는 전문가야.

    [핵심 규칙]
    1. 실제 회의 요약을 보고, 템플릿의 어느 위치(좌표ID)에 내용을 채워 넣어야 할지 판단해.
    2. **패턴 인식:** 표가 '한 사람당 1줄'인지, '한 사람당 2줄(병합)'인지 템플릿 내용을 보고 판단해.
       - 예: 1행에 '전혜나', 2행에 '향후계획', 3행에 '김성태'가 있다면 -> **"2행 1세트"** 구조다.
       - 이 경우, 실제 데이터도 1행(이름), 2행(계획), 3행(이름), 4행(계획) 순으로 건너뛰며 넣어야 한다.
    2. **샘플 데이터 무시**: 템플릿에 적혀 있는 이름(예: 홍길동, 김철수 등), 날짜, 내용은 모두 '예시'일 뿐이야. 의미를 해석해서 매칭하려 하지 마.
    3. **위치 기반 덮어쓰기**: - 표(Table)의 경우, 헤더(제목 줄)를 제외한 첫 번째 데이터 행부터 순서대로 새로운 데이터를 덮어씌워.
       - 예: 템플릿에 5명의 샘플이 있고, 실제 참석자가 2명이면 -> 1, 2번째 줄은 실제 데이터로 바꾸고.
    4. **잔여 데이터 삭제**:
       - 실제 데이터보다 템플릿의 행이 더 많다면, 남는 행의 내용을 반드시 비워야 해.
       - 방법: 해당 좌표의 `new_text`를 빈 문자열 `""`로 설정해.
       - 예: 3, 4, 5번째 줄의 샘플 데이터는 `""`로 교체하여 공란으로 만듦.
    5. **헤더 보존**: 표의 제목이나 항목명(예: "성명", "활동내역")은 건드리지 마.
    6. **누락 금지**: 참석 인원을 누락하지 마. 특히 표 안에 있는 팀원 이름, 활동 내역, 진행도 등은 줄 맞춰서 정확히 넣어줘.
    7. 내용이 없는 빈 칸이라도 문맥상 거기에 들어가야 한다면 해당 좌표를 지정해.
    8. 응답은 반드시 JSON 포맷으로 해.

    [입력 데이터 설명]
    - 템플릿 데이터는 '[좌표ID] 현재내용' 형식으로 되어 있어.
      예: "[T-0-R-1-C-2] 진행중" -> 0번 표, 1번 행, 2번 열에 "진행중"이 있다는 뜻.

    [작동 예시]
    Input(Template):
    [T-0-R-1-C-0] 샘플이름1 (참석자 A)
    [T-0-R-2-C-0] 샘플이름2 (참석자 B)
    [T-0-R-3-C-0] 샘플이름3 (참석자 C)
    
    Input(New Data): "참석자: 이영희" (1명)
    
    Output(JSON):
    {
        "updates": [
            {"id": "T-0-R-1-C-0", "new_text": "이영희"},
            {"id": "T-0-R-2-C-0", "new_text": ""},
            {"id": "T-0-R-3-C-0", "new_text": ""}
        ]
    }
    """
    
    user_prompt = f"""
    === [1] 템플릿 데이터 (좌표 포함) ===
    {template_text_with_coords}

    === [2] 실제 회의 요약 데이터 ===
    {new_summary_text}
    """

    print("🤖 AI가 문서를 분석하고 있습니다...")

    response = openai_client.chat.completions.create(
        model=AZURE_OPENAI_DEPLOYMENT_NAME,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        response_format={ "type": "json_object" }
    )

    return json.loads(response.choices[0].message.content)

# ==========================================
# 4. 문서 업데이트 (스타일 유지 + 잔여 문단 삭제)
# ==========================================
def update_docx_by_coordinates(template_path, output_path, updates):
    doc = Document(template_path)
    count = 0
    
    for item in updates:
        coord_id = item["id"]
        new_text = item["new_text"]
        
        try:
            target_paragraph = None
            
            # 표(Table) 처리
            if coord_id.startswith("T-"):
                parts = coord_id.split('-')
                t_idx, r_idx, c_idx = int(parts[1]), int(parts[3]), int(parts[5])
                cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
                
                # 셀 안의 첫 번째 문단을 타깃으로 설정. 셀 안에 문단이 여러 개라면 첫 번째만 남기고 삭제
                if len(cell.paragraphs) > 1:
                    # 뒤에서부터(역순으로) 제거 -> 인덱스 오류 방지
                    for i in range(len(cell.paragraphs) - 1, 0, -1):
                        p = cell.paragraphs[i]
                        # 문단 삭제 (XML 요소 제거)
                        p._element.getparent().remove(p._element)
                
                target_paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            # 문단(Paragraph) 처리
            elif coord_id.startswith("P-"):
                p_idx = int(coord_id.split('-')[1])
                target_paragraph = doc.paragraphs[p_idx]
            
            if target_paragraph:
                update_paragraph_with_style(target_paragraph, new_text)
                count += 1
                
        except IndexError:
            print(f"⚠️ 경고: 좌표 {coord_id}는 문서 범위를 벗어남 (Skipped)")
        except Exception as e:
            print(f"⚠️ 오류: {coord_id} 처리 실패: {e}")

    doc.save(output_path)
    print(f"✅ 문서 생성 완료. ({count}곳 변경됨): {output_path}")
    return output_path